import os
import sys
import time
import traceback

import pythoncom
import win32com.client
import pypandoc
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- CONFIGURATION ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Updated to use the new professional template
WORD_TEMPLATE = os.path.join(
    BASE_DIR, 
    "Professional_Template.docx"
)
TABLE_STYLE = 'Table Grid' 

# Dynamically pointing to the current directory
BASE_DIRECTORIES = [
    BASE_DIR
]

MD_PREFIX = "TechnicalDesign_"


def convert_docx_to_pdf_solid(docx_path, pdf_path):
    """
    Native converter that forces a new, isolated Microsoft Word process every time.
    This prevents memory leaks and COM object conflicts during batch processing.
    """
    pythoncom.CoInitialize()
    
    # DispatchEx creates a completely new and separate instance of Word
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    
    try:
        doc = word.Documents.Open(docx_path)
        # 17 is the internal Microsoft COM code for 'wdFormatPDF'
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close(0)  # Close the document without saving additional changes
    finally:
        word.Quit()
        pythoncom.CoUninitialize()


def convert_md_to_corporate_docs(input_md):
    """
    Reads a Markdown file, converts it to DOCX via Pandoc using a corporate template,
    applies specific styling to paragraphs and tables, and exports the result to PDF.
    """
    input_md = os.path.abspath(str(input_md))
    base_name = os.path.splitext(input_md)[0]
    
    output_docx = os.path.abspath(base_name + ".docx")
    final_pdf = os.path.abspath(base_name + ".pdf")
    
    if not os.path.exists(WORD_TEMPLATE):
        print(f"Error: Word template not found at:\n{WORD_TEMPLATE}")
        return

    try:
        print("Step 1: Converting Markdown to DOCX via Pandoc...")
        
        # Pandoc requires forward slashes for template paths
        pandoc_template = WORD_TEMPLATE.replace('\\', '/')
        extra_args = [f'--reference-doc={pandoc_template}']
        
        with open(input_md, 'r', encoding='utf-8') as file:
            md_content = file.read()
            
        pypandoc.convert_text(
            md_content, 
            'docx', 
            format='markdown-yaml_metadata_block',
            outputfile=output_docx, 
            extra_args=extra_args
        )

        print("Step 2: Formatting tables and paragraph styles...")
        doc = Document(output_docx)
        
        # --- PARAGRAPH & CODE BLOCK FORMATTING ---
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name.lower()
            
            # 1. Format Major Headers (Heading 1, Heading 2, Heading 3)
            if any(h in style_name for h in ['heading 1', 'heading 2', 'heading 3']):
                paragraph.paragraph_format.space_before = Pt(24) # Extra air before new section
                paragraph.paragraph_format.space_after = Pt(12)  # Air between title and text
                paragraph.paragraph_format.keep_with_next = True # Keep title attached to paragraph
            
            # 2. Format Measure Headers (Heading 4 in Markdown becomes Heading 4 in Word)
            elif 'heading 4' in style_name:
                paragraph.paragraph_format.space_before = Pt(28) # Huge air before new measure
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.keep_with_next = True
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(8, 49, 42) # Corporate Dark Green
                    run.font.size = Pt(13)
                    run.bold = True
                    
            # 3. Add breathing room to Bullet Points
            elif 'list paragraph' in style_name or 'list' in style_name:
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 4. Advanced Styling for DAX Code Blocks (The "Web-Look" Hack)
            elif 'code' in style_name or 'source' in style_name:
                # Left alignment and spacing
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(8)
                paragraph.paragraph_format.space_after = Pt(24) # Push next content down heavily
                paragraph.paragraph_format.line_spacing = 1.15
                
                # Add a slight visual indent from the main text
                paragraph.paragraph_format.left_indent = Pt(10)
                paragraph.paragraph_format.right_indent = Pt(10)
                
                # Prevent code blocks from breaking across pages
                paragraph.paragraph_format.keep_lines_together = True
                
                # GitHub Dark Theme Colors
                bg_color = "0D1117"  # Very dark, elegant slate/navy
                
                # Add Dark Background Shading via XML
                shd_xml = r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color)
                shading_element = parse_xml(shd_xml)
                paragraph._p.get_or_add_pPr().append(shading_element)
                
                # THE HACK: Add a border of the EXACT SAME COLOR to act as internal padding
                border_xml = r'<w:pBdr {}><w:top w:val="single" w:sz="12" w:space="8" w:color="{}"/><w:left w:val="single" w:sz="12" w:space="8" w:color="{}"/><w:bottom w:val="single" w:sz="12" w:space="8" w:color="{}"/><w:right w:val="single" w:sz="12" w:space="8" w:color="{}"/></w:pBdr>'.format(nsdecls('w'), bg_color, bg_color, bg_color, bg_color)
                border_element = parse_xml(border_xml)
                paragraph._p.get_or_add_pPr().append(border_element)

                # Format the text inside the code block
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(201, 209, 217)  # Soft pearl white/gray
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)  # Standard developer font size
            
            # 5. Justify standard text (excluding headers, titles, and lists)
            elif not any(excluded in style_name for excluded in ['heading', 'title', 'list']):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- TABLE FORMATTING ---
        for table in doc.tables:
            table.style = TABLE_STYLE
            table.autofit = False 
            table.allow_autofit = False
            
            num_cols = len(table.columns)
            
            # --- NATIVE COLUMN WIDTHS (Based on 6.5 inches total usable page width) ---
            if num_cols == 3:
                widths = [Inches(1.5), Inches(3.5), Inches(1.5)] 
            elif num_cols == 4:
                widths = [Inches(1.5), Inches(1.5), Inches(2.0), Inches(1.5)] 
            elif num_cols == 5:
                widths = [Inches(1.5), Inches(1.0), Inches(1.0), Inches(1.0), Inches(2.0)] 
            else:
                widths = [Inches(6.5 / num_cols)] * num_cols 
            
            # Apply widths to column definitions
            for j, col in enumerate(table.columns):
                col.width = widths[j]
            
            for i, row in enumerate(table.rows):
                # Prevent rows from splitting across pages
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                cant_split = parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w')))
                trPr.append(cant_split)
                
                for j, cell in enumerate(row.cells):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.width = widths[j] # Apply explicit width to every cell
                    
                    # [NUEVO] Convertir el token [NL] en saltos de línea reales de Word
                    if '[NL]' in cell.text:
                        cell.text = cell.text.replace('[NL]', '\n')
                    
                    if i == 0:
                        # Header row styling (Corporate Dark Green color)
                        shading_element = parse_xml(r'<w:shd {} w:fill="08312A"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading_element)
                        
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        # Data rows alternating background color (Light Gray)
                        if i % 2 == 0:
                            shading_element = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shading_element)
                        
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
        doc.save(output_docx)

        # Pause to prevent file-locking issues with cloud synchronization tools
        print("Waiting 3 seconds for cloud synchronization release...")
        time.sleep(3)

        print("Step 3: Exporting to PDF (Isolated COM Mode)...")
        convert_docx_to_pdf_solid(output_docx, final_pdf)
        
        print("Success: Corporate documents generated in the source directory.")
        
    except Exception as e:
        print(f"Error processing file: {e}")
        print("--- Technical Details ---")
        traceback.print_exc()
        print("-------------------------")


def search_and_convert_all():
    """
    Scans the configured base directories for Markdown files matching 
    the prefix and triggers the conversion process.
    """
    print("Initiating automated search for Markdown documentation...")
    files_found = 0
    
    for base_dir in BASE_DIRECTORIES:
        if not os.path.exists(base_dir):
            print(f"Warning: Cannot access directory: {base_dir}")
            continue
            
        print(f"Scanning recursively: {base_dir}")
        
        # Ignore .venv and .git directories
        for root, dirs, files in os.walk(base_dir):
            if '.venv' in dirs: dirs.remove('.venv')
            if '.git' in dirs: dirs.remove('.git')
            
            for file in files:
                if file.startswith(MD_PREFIX) and file.endswith(".md"):
                    files_found += 1
                    filepath = os.path.join(root, file)
                    
                    print("\n" + "-" * 50)
                    print(f"Processing [{files_found}]: {file}")
                    convert_md_to_corporate_docs(filepath)
                    
    if files_found == 0:
        print("\nNo pending Markdown files found.")
    else:
        print(f"\nBatch process completed. {files_found} documents processed successfully.")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        target_md = sys.argv[1]
        print("\n" + "-" * 50)
        print(f"Processing single file: {os.path.basename(target_md)}")
        convert_md_to_corporate_docs(target_md)
    else:
        search_and_convert_all()