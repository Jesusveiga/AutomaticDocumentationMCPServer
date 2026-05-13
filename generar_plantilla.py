import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_template():
    doc = Document()

    # --- PALETA DE COLORES CORPORATIVA ---
    PRIMARY_COLOR = RGBColor(8, 49, 42) # Verde oscuro elegante (SDG/Corporate)
    TEXT_COLOR = RGBColor(60, 60, 60)   # Gris carbón para no fatigar la vista
    WHITE = RGBColor(255, 255, 255)

    # 1. Configurar Estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = TEXT_COLOR
    style_normal.paragraph_format.space_after = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.15

    # 2. Configurar Estilo de Título Principal (Portada/Cabecera)
    style_title = doc.styles['Title']
    style_title.font.name = 'Segoe UI'
    style_title.font.size = Pt(28)
    style_title.font.bold = True
    style_title.font.color.rgb = PRIMARY_COLOR
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_title.paragraph_format.space_after = Pt(30)

    # 3. Configurar Estilo de Título 1 (Secciones)
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Segoe UI'
    style_h1.font.size = Pt(16)
    style_h1.font.bold = True
    style_h1.font.color.rgb = PRIMARY_COLOR
    style_h1.paragraph_format.space_before = Pt(24)
    style_h1.paragraph_format.space_after = Pt(12)

    # --- PIE DE PÁGINA DISTINTIVO ---
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Automated Document Generation via MCP | SDG Group"
    footer_para.style.font.name = 'Segoe UI'
    footer_para.style.font.size = Pt(9)
    footer_para.style.font.color.rgb = RGBColor(150, 150, 150)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Función auxiliar para pintar viñetas rápidas
    def add_bullet(key, value):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(key).bold = True
        p.add_run(value)

    # ==========================================
    # --- INYECCIÓN DEL CONTENIDO ESTRICTO ---
    # ==========================================
    
    doc.add_paragraph('Technical Design & Specifications', style='Title')

    # Sección 1
    doc.add_heading('1. Introduction & Executive Summary', level=1)
    doc.add_paragraph('This document outlines the technical architecture, data model, and measure calculations for the Power BI semantic model. It serves as a comprehensive guide for data engineers, BI developers, and business stakeholders to ensure alignment on metric definitions and data governance.')
    add_bullet('Report Name: ', '[Demo] Global GFE Balanced Scorecard')
    add_bullet('Business Owner: ', 'Sales Operations Dept.')
    add_bullet('Technical Owner: ', 'Jesús Veiga Morandeira / SDG Group')
    add_bullet('Target Audience: ', 'Executive Board, Regional Managers')

    # Sección 2
    doc.add_heading('2. Data Sources & Architecture', level=1)
    doc.add_paragraph('Brief description of where the data originates before entering the Power BI model.')
    add_bullet('Source Systems: ', 'Contoso Azure SQL Database / Apptio / Flat Files')
    add_bullet('Data Integration: ', 'Import Mode')
    add_bullet('Refresh Schedule: ', 'Daily at 06:00 AM UTC')

    # Sección 3
    doc.add_heading('3. Security & Governance', level=1)
    doc.add_paragraph('Defines the access controls applied to the semantic model.')
    add_bullet('Workspace: ', 'PRD_Contoso_Global_Sales')
    add_bullet('Row-Level Security (RLS): ', 'Yes - Filtered by RegionManagerID')
    add_bullet('Sensitivity Label: ', 'Highly Confidential / Internal')

    # Sección 4 (ZONA CERO)
    doc.add_heading('4. Semantic Model & Report Design', level=1)
    p_auto = doc.add_paragraph()
    run_auto = p_auto.add_run('[Auto-generated via MCP Pipeline. The AI Agent will inject sections 4.2 Data Model and 4.3 Report Design directly below this point.]')
    run_auto.italic = True
    run_auto.font.color.rgb = RGBColor(150, 150, 150)

    # Sección 5
    doc.add_heading('5. Change Log (Version Control)', level=1)
    
    # Crear y estilizar la tabla de forma corporativa
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    headers = ['Version', 'Date', 'Author', 'Description of Changes']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        
        # Fondo verde corporativo para la cabecera
        shading_elm = parse_xml(r'<w:shd {} w:fill="08312A"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Texto en blanco y negrita
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                
    row_cells = table.rows[1].cells
    row_cells[0].text = '1.0'
    row_cells[1].text = 'April 2026'
    row_cells[2].text = 'Jesús Veiga'
    row_cells[3].text = 'Initial automated documentation release via MCP'

    # Guardar documento
    file_name = 'Professional_Template.docx'
    doc.save(file_name)
    print(f"\n¡Éxito! Plantilla corporativa generada en:\n{os.path.abspath(file_name)}")

if __name__ == "__main__":
    create_template()